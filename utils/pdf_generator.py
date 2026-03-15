"""Resume PDF Generator - Matches master resume format exactly.

Generates a .docx using python-docx with the same tight formatting
as the master resume, then converts to PDF via:
1. docx2pdf (uses Microsoft Word on Mac / Windows)
2. LibreOffice (Linux fallback)
3. If neither available, leaves the .docx for manual conversion
"""

import re
import subprocess
import shutil
from pathlib import Path


def markdown_to_pdf(md_text: str, output_path: str, name: str = None) -> str:
    """Convert tailored resume markdown to a professional PDF.

    Returns path to the generated PDF (or .docx if no converter available).
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    parsed = _parse_resume_sections(md_text)
    docx_path = output_path.with_suffix(".docx")
    _build_docx(parsed, docx_path)

    # Try converting to PDF
    pdf_path = output_path.with_suffix(".pdf")

    # Method 1: LibreOffice (most reliable on Mac/Linux)
    soffice_cmd = None
    if shutil.which("libreoffice"):
        soffice_cmd = "libreoffice"
    elif shutil.which("soffice"):
        soffice_cmd = "soffice"
    elif Path("/Applications/LibreOffice.app/Contents/MacOS/soffice").exists():
        soffice_cmd = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

    if soffice_cmd:
        try:
            result = subprocess.run(
                [soffice_cmd, "--headless", "--convert-to", "pdf",
                 str(docx_path), "--outdir", str(output_path.parent)],
                capture_output=True, timeout=60,
            )
            if pdf_path.exists():
                print(f"   PDF generated: {pdf_path}")
                return str(pdf_path)
        except Exception:
            pass

    # Method 2: docx2pdf (uses Word on Mac/Windows — can be flaky)
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            print(f"   PDF generated: {pdf_path}")
            return str(pdf_path)
    except Exception:
        pass

    # No converter - return docx
    print(f"   DOCX generated: {docx_path}")
    print(f"   To get PDF, install: pip install docx2pdf")
    return str(docx_path)


def _parse_resume_sections(md_text: str) -> dict:
    """Parse markdown resume into structured sections."""
    sections = {
        "header": {"name": "", "contact": ""},
        "education": [],
        "skills": [],
        "experience": [],
        "research": [],
        "projects": [],
    }

    lines = md_text.strip().split("\n")
    current_section = None
    current_entry = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Name
        if stripped.startswith("# ") and not stripped.startswith("## "):
            sections["header"]["name"] = stripped[2:].strip()
            continue

        # Contact line
        if not stripped.startswith("#") and ("@" in stripped or "|" in stripped) and not current_section:
            sections["header"]["contact"] = stripped
            # If we have a contact line but no name yet, the previous non-empty line was probably the name
            if not sections["header"]["name"] and len(lines) > 1:
                for prev_line in lines[:lines.index(line)]:
                    prev = prev_line.strip()
                    if prev and not prev.startswith("#") and "@" not in prev and "|" not in prev:
                        sections["header"]["name"] = prev
                        break
            continue

        # Name without # prefix — first non-empty line before contact info
        if not current_section and not sections["header"]["name"] and not sections["header"]["contact"]:
            if not stripped.startswith("#") and "@" not in stripped and "|" not in stripped:
                sections["header"]["name"] = stripped
                continue

        # Section headers
        if stripped.startswith("## "):
            section_name = stripped[3:].strip().lower()
            if "education" in section_name:
                current_section = "education"
            elif "skill" in section_name:
                current_section = "skills"
            elif "experience" in section_name:
                current_section = "experience"
            elif "research" in section_name:
                current_section = "research"
            elif "project" in section_name:
                current_section = "projects"
            current_entry = None
            continue

        # Entry headers
        if stripped.startswith("### "):
            entry_text = stripped[4:].strip()
            current_entry = {"title": entry_text, "subtitle": "", "bullets": []}
            if current_section and current_section in sections:
                sections[current_section].append(current_entry)
            continue

        # Bullets
        if (stripped.startswith("- ") or stripped.startswith("* ")):
            bullet = stripped[2:].strip()
            if current_entry:
                current_entry["bullets"].append(bullet)
            elif current_section == "skills":
                sections["skills"].append(bullet)
            continue

        # Non-bullet text in skills section (plain lines like "Languages: Python, SQL...")
        if current_section == "skills" and not current_entry and ":" in stripped:
            sections["skills"].append(stripped)
            continue

        # Subtitle / other text under an entry
        if current_entry:
            if current_entry["subtitle"]:
                current_entry["subtitle"] += "\n" + stripped
            else:
                current_entry["subtitle"] = stripped

    return sections


def _build_docx(parsed: dict, output_path: Path):
    """Build a .docx matching the master resume format exactly."""
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup - tight margins matching master resume
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Twips(540)   # ~0.375in
    section.bottom_margin = Twips(540)
    section.left_margin = Twips(620)  # ~0.43in
    section.right_margin = Twips(620)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # -- Name --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(parsed["header"]["name"])
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    # -- Contact --
    if parsed["header"]["contact"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(parsed["header"]["contact"])
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        # Bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "1", qn("w:color"): "000000",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_entry_header(title_text):
        """Parse 'Company | Role (Date)' and format with bold company, right-aligned date."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)

        # Extract date from parentheses
        date_match = re.search(r'\(([^)]+)\)\s*$', title_text)
        date_str = ""
        if date_match:
            date_str = date_match.group(1)
            title_text = title_text[:date_match.start()].strip()

        # Split on | for company | role
        if "|" in title_text:
            parts = title_text.split("|", 1)
            run = p.add_run(parts[0].strip())
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            run = p.add_run(" | " + parts[1].strip())
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
        else:
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

        # Add date with tab stop
        if date_str:
            # Set right tab stop at page width minus margins
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.64), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            run = p.add_run(f"\t{date_str}")
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            run.italic = True

        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.font.name = "Times New Roman"
        run.italic = True

    def add_bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        run = p.add_run("\u2022 " + text)
        run.font.size = Pt(8.5)
        run.font.name = "Times New Roman"

    def add_skill_line(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if ":" in text:
            label, value = text.split(":", 1)
            run = p.add_run(label.strip() + ":")
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            run = p.add_run(value)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
        else:
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

    def add_entries(entries):
        for entry in entries:
            add_entry_header(entry["title"])
            if entry.get("subtitle"):
                for sub_line in entry["subtitle"].split("\n"):
                    add_subtitle(sub_line)
            for bullet in entry.get("bullets", []):
                add_bullet(bullet)

    # -- Build sections --
    if parsed["education"]:
        add_section_header("Education")
        add_entries(parsed["education"])

    if parsed["skills"]:
        add_section_header("Technical Skills")
        for line in parsed["skills"]:
            add_skill_line(line)

    if parsed["experience"]:
        add_section_header("Experience")
        add_entries(parsed["experience"])

    if parsed["research"]:
        add_section_header("Research")
        add_entries(parsed["research"])

    if parsed["projects"]:
        add_section_header("Projects")
        add_entries(parsed["projects"])

    doc.save(str(output_path))